#!/usr/bin/env python3
"""
Generate AXIOM-Mobile research paper as a publication-quality .docx file.
Uses python-docx for document creation with academic formatting.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "AXIOM_Mobile_Paper.docx")

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="000000", size="4"):
    """Set all borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_table_with_borders(doc, headers, rows, caption=None, caption_number=None):
    """Add a formatted table with borders, header shading, and optional caption."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        set_cell_shading(cell, "D9D9D9")
        set_cell_borders(cell)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            set_cell_borders(cell)

    # Caption below table
    if caption and caption_number is not None:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(10)
        run_label = cap.add_run(f"Table {caption_number}: ")
        run_label.bold = True
        run_label.italic = True
        run_label.font.size = Pt(9)
        run_label.font.name = "Times New Roman"
        run_text = cap.add_run(caption)
        run_text.italic = True
        run_text.font.size = Pt(9)
        run_text.font.name = "Times New Roman"

    return table


def add_body(doc, text):
    """Add a body-text paragraph (10pt TNR, justified, 6pt spacing)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    return p


def add_heading_numbered(doc, number, title, level=1):
    """Add a numbered section/subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}  {title}")
    run.bold = True
    if level == 1:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    return p


def add_page_numbers(doc):
    """Add page numbers to the footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)
    for r in [run, run2, run3]:
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"


# ── Main document generation ─────────────────────────────────────────────────

def generate():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(12)

    # ── Title ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run(
        "AXIOM-Mobile: Measuring Minimal Training Data Requirements\n"
        "for On-Device Domain Reasoning Under Mobile Constraints"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # ── Authors ──────────────────────────────────────────────────────────
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(2)
    run = author_p.add_run("Annie Boltwood, Mahim Chaudhary, Ariel Tyson")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    affil_p = doc.add_paragraph()
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_p.paragraph_format.space_after = Pt(16)
    run = affil_p.add_run("Simon Fraser University, School of Computing Science")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.italic = True

    # ── Abstract ─────────────────────────────────────────────────────────
    abs_p = doc.add_paragraph()
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.left_indent = Inches(0.5)
    abs_p.paragraph_format.right_indent = Inches(0.5)
    abs_p.paragraph_format.space_after = Pt(12)
    abs_p.paragraph_format.line_spacing = Pt(12)
    label = abs_p.add_run("Abstract\u2014")
    label.bold = True
    label.italic = True
    label.font.size = Pt(10)
    label.font.name = "Times New Roman"
    abs_text = abs_p.add_run(
        "We present AXIOM-Mobile, an integrated system and experimental framework for measuring "
        "the minimal training set size (k*) required for effective visual question answering on "
        "mobile devices under strict quality, latency, energy, and memory constraints. We report "
        "results across two experimental regimes: v0, comprising 52 examples over 24 answer classes "
        "yielding 10% test exact match (EM), and v1, comprising 452 examples over 128 answer classes "
        "yielding 27.5% test EM. Physical-device profiling on an iPhone 15 Pro Max (A17 Pro) "
        "achieves p50 latency of 14.0\u201314.5 ms, which is 28\u00d7 below the 400 ms threshold. "
        "While the 70% EM quality target remains unmet, dataset scaling is confirmed as the "
        "primary lever for improvement. We employ an honest status tracking system that "
        "transparently reports metric shortfalls alongside infrastructure achievements."
    )
    abs_text.font.size = Pt(10)
    abs_text.font.name = "Times New Roman"

    # ── Section 1: Introduction ──────────────────────────────────────────
    add_heading_numbered(doc, "1", "Introduction")

    add_body(doc,
        "Mobile devices have become the primary computing platform for billions of users worldwide, "
        "yet deploying machine learning models that perform domain-specific reasoning on-device "
        "remains a fundamental challenge. The tension is clear: achieving high-quality predictions "
        "typically demands large training datasets and expressive model architectures, while "
        "practical mobile deployment requires small, fast models that operate within strict "
        "resource budgets. This tension is particularly acute for visual question answering (VQA) "
        "tasks, where models must jointly reason over image content and natural language queries "
        "to produce correct responses."
    )

    add_body(doc,
        "In this paper, we investigate a precise formulation of this tension through the following "
        "research question: What is k*\u2014the minimal training set size at which an on-device "
        "VQA model achieves effective domain reasoning under mobile constraints? We define "
        "\"effective\" through six measurable thresholds that span quality and efficiency: "
        "exact match accuracy EM \u2265 70%, median inference latency p50 \u2264 400 ms, "
        "95th-percentile latency p95 \u2264 600 ms, energy consumption < 5% battery per hour, "
        "peak memory usage < 500 MB, and model binary size < 100 MB. These thresholds are "
        "grounded in user-experience research and mobile platform guidelines, ensuring that "
        "our definition of effectiveness is both rigorous and practically relevant."
    )

    add_body(doc,
        "To address this question, we make five contributions. First, we develop an end-to-end "
        "iOS and Python testbed with real Core ML inference on physical hardware, enabling "
        "reproducible benchmarking that reflects actual deployment conditions. Second, we "
        "implement and compare three data selection strategies\u2014random, uncertainty-based, and "
        "diversity-based\u2014with bootstrap confidence intervals to quantify differences under "
        "limited data regimes. Third, we build a statistical analysis pipeline that incorporates "
        "an honest status vocabulary, transparently reporting when metrics are met, unmet, or "
        "unavailable. Fourth, we construct a scaled dataset (v2) that achieves a 2.75\u00d7 "
        "improvement in test exact match over the initial dataset. Fifth, we perform physical-device "
        "latency profiling demonstrating 14 ms median inference\u201428\u00d7 below the quality-of-service "
        "threshold. Our work builds upon foundations in active learning and core-set methods [1][2], "
        "visual question answering [3], and neural scaling laws [4]."
    )

    # ── Section 2: Related Work ──────────────────────────────────────────
    add_heading_numbered(doc, "2", "Related Work")

    add_body(doc,
        "Data-efficient learning has been extensively studied through the lens of active learning "
        "and core-set selection. Sener and Savarese [1] formulate active learning as a core-set "
        "selection problem, seeking to choose a subset of training examples that approximates the "
        "full dataset's geometric structure. Ash et al. [2] extend this direction by combining "
        "uncertainty and diversity through gradient embeddings, achieving strong performance with "
        "limited labeled data. Settles [8] provides a comprehensive survey of active learning "
        "strategies. Our work applies these principles to mobile VQA, where device constraints "
        "impose requirements beyond accuracy alone\u2014latency, energy, and memory budgets must "
        "simultaneously be satisfied."
    )

    add_body(doc,
        "On-device machine learning has gained considerable attention with frameworks such as "
        "Core ML [5] and TensorFlow Lite enabling neural network inference on mobile hardware. "
        "However, most published benchmarks report only top-line accuracy or measure performance "
        "on desktop GPUs, leaving a gap in understanding real-world mobile behavior. Ignatov et al. "
        "[6] address this partially through their AI Benchmark suite, which profiles deep learning "
        "operations on smartphone processors. Howard et al. [9] develop MobileNetV3, demonstrating "
        "that architecture search can yield efficient models for mobile platforms. We follow the "
        "philosophy of profiling on target hardware, extending it to VQA workloads with "
        "per-inference latency distributions rather than aggregate throughput."
    )

    add_body(doc,
        "Visual question answering, introduced by Antol et al. [3] through the VQA dataset, "
        "has catalyzed significant progress in multimodal reasoning. Recent work on visual "
        "instruction tuning [10] has pushed VQA capabilities substantially forward, though "
        "these approaches assume server-side inference with large pretrained backbones. Mobile "
        "VQA on domain-specific content\u2014such as app screenshots\u2014remains underexplored. "
        "Our setting differs from standard VQA in that questions target functional UI elements "
        "rather than natural scenes, and models must operate within strict mobile resource envelopes."
    )

    add_body(doc,
        "Scaling laws for neural networks, established by Hestness et al. [7] and Kaplan et al. "
        "[4], describe power-law relationships between dataset size, model size, compute budget, "
        "and performance. These laws have proven remarkably consistent across language and vision "
        "domains. We attempt to fit power-law models to our learning curves but report honestly "
        "when the data is too sparse\u2014with only six budget levels and three seeds, our fits "
        "yield low R\u00b2 values that do not support reliable extrapolation."
    )

    # ── Section 3: System Overview ───────────────────────────────────────
    add_heading_numbered(doc, "3", "System Overview")

    add_body(doc,
        "AXIOM-Mobile consists of two tightly integrated components: an iOS/macOS application "
        "for on-device inference and benchmarking, and a Python pipeline for dataset management, "
        "model training, and statistical analysis. Together, these components form a closed-loop "
        "system in which models are trained in Python, exported to Core ML format, deployed to "
        "the iOS application, profiled on physical hardware, and the resulting metrics are "
        "ingested back into the analysis pipeline."
    )

    add_heading_numbered(doc, "3.1", "iOS/macOS Application", level=2)

    add_body(doc,
        "The application is built with SwiftUI and targets iOS 17+ and macOS 14+. Its core "
        "functionality centers on a testbed interface that supports screenshot import, free-form "
        "question input, and model selection through a dedicated picker. Inference is performed "
        "through CoreMLInferenceService, which loads Core ML model packages and executes "
        "predictions using the Neural Engine or GPU as determined by the system scheduler. "
        "Each model is accompanied by a metadata sidecar file that specifies input dimensions, "
        "normalization parameters, vocabulary mappings, and confidence thresholds, enabling "
        "the application to configure itself automatically for each model variant."
    )

    add_body(doc,
        "The application supports a benchmark mode that executes between 1 and 50 inference "
        "iterations on a given input, recording per-iteration latency in milliseconds. An "
        "auto-benchmark mode extends this by cycling through a BenchmarkInputProvider that "
        "supplies deterministic input pairs, ensuring reproducibility across profiling sessions. "
        "Results are exported as CSV files with accompanying metadata, facilitating direct "
        "ingestion by the Python analysis pipeline."
    )

    add_heading_numbered(doc, "3.2", "Python Pipeline", level=2)

    add_body(doc,
        "The Python pipeline manages the complete experimental workflow. Dataset management "
        "operates on JSONL manifests that associate screenshot paths, questions, and ground-truth "
        "answers with split assignments (pool, validation, test). A model harness provides a "
        "unified interface for training, prediction, and Core ML export, with an accuracy gate "
        "that verifies the exported model suffers no more than a 3% exact match drop relative "
        "to the PyTorch source. Three selection strategies\u2014random, uncertainty, and diversity"
        "\u2014implement different approaches to choosing training examples from the pool under "
        "a given budget constraint."
    )

    add_body(doc,
        "The statistical analysis module computes bootstrap confidence intervals for all "
        "performance metrics, performs paired comparisons between strategies at each budget "
        "level, attempts power-law fits to learning curves, and constructs Pareto views "
        "that jointly consider quality and efficiency metrics. A device-profile ingestion "
        "module parses CSV exports from the iOS application, computes latency percentiles, "
        "and evaluates each session against the six effectiveness thresholds."
    )

    # ── Section 4: Dataset ───────────────────────────────────────────────
    add_heading_numbered(doc, "4", "Dataset")

    add_heading_numbered(doc, "4.1", "Dataset v1", level=2)

    add_body(doc,
        "The initial dataset (v1) comprises 52 screenshot-question-answer triples drawn from "
        "real iOS applications. Screenshots capture diverse UI states including home screens, "
        "settings panels, notification centers, and in-app interfaces. Questions target "
        "factual properties of the displayed content, such as \"What time is shown?\" or "
        "\"How many unread notifications are visible?\" The dataset is partitioned into a pool "
        "of 37 examples for training, 5 for validation, and 10 for testing. Answer normalization "
        "yields 24 distinct classes. This dataset supports all learning curve experiments reported "
        "in Section 7."
    )

    add_heading_numbered(doc, "4.2", "Dataset v2", level=2)

    add_body(doc,
        "To investigate the effect of dataset scale on model quality, we construct an expanded "
        "dataset (v2) containing 452 examples derived from 152 screenshots\u201452 manually captured "
        "and 100 auto-generated through programmatic screen composition. Answer normalization "
        "produces 128 distinct classes, a 5.3\u00d7 increase over v1. The dataset is split into a "
        "pool of 382 examples, 30 for validation, and 40 for testing\u2014representing an 8.7\u00d7 "
        "increase in pool size. This expanded dataset enables training of the v1 model variant "
        "and provides preliminary evidence for the role of dataset scale in improving quality."
    )

    add_heading_numbered(doc, "4.3", "Limitations", level=2)

    add_body(doc,
        "Several limitations constrain the conclusions that can be drawn from these datasets. "
        "Both versions remain below the 500-example target that we originally projected for "
        "reliable scaling analysis. All annotations are produced by a single annotator; a "
        "dual-annotator protocol with Cohen's kappa inter-rater reliability is planned but "
        "not yet implemented. Of the 152 screenshots in v2, 100 are auto-generated rather "
        "than captured from genuine device usage, which may introduce distributional differences. "
        "Finally, no selection-strategy sweep has been conducted on dataset v2, meaning that "
        "scaling-law fits and strategy comparisons are available only for the smaller v1 dataset."
    )

    # ── Section 5: Models ────────────────────────────────────────────────
    add_heading_numbered(doc, "5", "Models")

    add_heading_numbered(doc, "5.1", "Heuristic Baseline (question_lookup_v0)", level=2)

    add_body(doc,
        "As a lower bound on performance, we implement a zero-dependency heuristic baseline "
        "that memorizes the most common answer for each normalized question string observed "
        "in the training pool. This model requires no gradient-based optimization and produces "
        "deterministic outputs. On the training pool it achieves 72.97% exact match, reflecting "
        "pure memorization. On the held-out test set it achieves 10% exact match, demonstrating "
        "that question-level memorization does not generalize to unseen screenshots. This baseline "
        "establishes the floor against which learned models must improve."
    )

    add_heading_numbered(doc, "5.2", "Trainable Baseline (tiny_multimodal_v0)", level=2)

    add_body(doc,
        "The trainable baseline is a compact multimodal architecture with approximately 40,000 "
        "parameters. The image encoder is a three-layer convolutional neural network that "
        "processes 128\u00d7128 RGB inputs and produces a 64-dimensional feature vector. The text "
        "encoder uses character-level embeddings over an ASCII vocabulary of 128 tokens, with "
        "mean pooling yielding a 64-dimensional representation. Image and text features are "
        "concatenated into a 128-dimensional joint representation, which is passed through a "
        "linear classification head with 24 output classes. The model achieves 16.22% pool EM "
        "and 10% test EM on dataset v1. When exported to Core ML, the model package is 96 KB, "
        "well within the 100 MB size threshold."
    )

    add_heading_numbered(doc, "5.3", "Retrained Model (tiny_multimodal_v1)", level=2)

    add_body(doc,
        "The v1 model retains the same architectural template as v0 but is retrained on dataset "
        "v2 with an expanded output head of 128 classes, resulting in approximately 47,000 "
        "parameters. Training employs class-weighted cross-entropy loss with inverse-frequency "
        "weighting capped at 10\u00d7 to prevent extreme gradients from rare classes, and runs for "
        "40 epochs. The model achieves 30.9% pool EM, 26.7% validation EM, and 27.5% test EM, "
        "representing a 2.75\u00d7 improvement in test accuracy over the v0 baseline. The Core ML "
        "export passes the accuracy gate with 0% drop, and a calibrated confidence threshold "
        "of 0.45 is selected based on validation-set analysis. Table 1 summarizes the comparison "
        "between the two model versions."
    )

    add_table_with_borders(doc,
        headers=["Metric", "v0 (dataset v1)", "v1 (dataset v2)"],
        rows=[
            ["Pool EM", "16.22%", "30.9%"],
            ["Val EM", "0.00%", "26.7%"],
            ["Test EM", "10.00%", "27.5%"],
            ["Parameters", "~40K", "~47K"],
            ["Output Classes", "24", "128"],
        ],
        caption="Model comparison between tiny_multimodal v0 and v1.",
        caption_number=1
    )

    # ── Section 6: Selection Strategies and Sweep Design ─────────────────
    add_heading_numbered(doc, "6", "Selection Strategies and Sweep Design")

    add_body(doc,
        "We implement three data selection strategies to investigate how the choice of training "
        "examples affects learning efficiency under budget constraints. The random strategy "
        "samples uniformly from the unlabeled pool without replacement. The uncertainty strategy "
        "selects examples for which the current model exhibits the highest predictive entropy, "
        "prioritizing inputs where the model is least confident. The diversity strategy employs "
        "k-center greedy selection in feature space, choosing examples that maximize the minimum "
        "distance to already-selected points and thereby encouraging broad coverage of the input "
        "distribution. A fourth strategy based on knowledge-graph guidance was planned but is "
        "blocked pending construction of a suitable domain ontology."
    )

    add_body(doc,
        "The experimental sweep is defined as a grid over three strategies, six budget levels "
        "(5, 10, 15, 20, 25, and 37 examples), and three random seeds, yielding 54 total runs. "
        "Each run trains a fresh model instance on the selected subset, evaluates on the held-out "
        "test set, and records exact match accuracy. An important caveat applies: all sweep "
        "experiments are conducted exclusively on dataset v1, meaning that results reflect the "
        "52-example regime and may not extrapolate to the larger v2 dataset."
    )

    # ── Section 7: Experiments ───────────────────────────────────────────
    add_heading_numbered(doc, "7", "Experiments")

    add_heading_numbered(doc, "7.1", "Learning Curves", level=2)

    add_body(doc,
        "Table 2 reports test exact match accuracy across budget levels for each selection "
        "strategy, averaged over three seeds. The results reveal a strikingly flat landscape: "
        "no strategy achieves above 13.3% at any budget, and all converge to 10% at the "
        "maximum budget of 37 (the full pool). This pattern indicates that the v0 model "
        "architecture, with only 40K parameters and no pretrained backbone, lacks sufficient "
        "capacity to learn robust visual features from the available data."
    )

    add_table_with_borders(doc,
        headers=["Budget", "Random", "Diversity", "Uncertainty"],
        rows=[
            ["5", "0.100", "0.067", "0.000"],
            ["10", "0.033", "0.067", "0.000"],
            ["15", "0.067", "0.133", "0.000"],
            ["20", "0.067", "0.133", "0.000"],
            ["25", "0.067", "0.067", "0.000"],
            ["37", "0.100", "0.100", "0.100"],
        ],
        caption="Test exact match by budget and selection strategy (dataset v1, 3 seeds).",
        caption_number=2
    )

    add_body(doc,
        "We attempt power-law fits of the form EM = a \u00b7 k^b to each strategy's learning "
        "curve. The diversity strategy yields R\u00b2 = 0.17, the random strategy R\u00b2 = 0.02, "
        "and the uncertainty strategy produces a degenerate fit due to zero accuracy at all "
        "budgets below 37. These R\u00b2 values are too low to support extrapolation or to draw "
        "conclusions about scaling behavior. We attribute this to the combination of an "
        "under-parameterized model and an insufficient number of budget levels and seeds."
    )

    add_heading_numbered(doc, "7.2", "Strategy Comparisons", level=2)

    add_body(doc,
        "To test whether any strategy significantly outperforms the others, we conduct paired "
        "bootstrap comparisons at budget level 37 (full pool). All pairwise differences in test "
        "exact match are 0.000 with 95% confidence intervals of [0.000, 0.000]. This result "
        "is unsurprising given that all strategies select the entire pool at budget 37, yielding "
        "identical training sets. At lower budgets, the small number of seeds (three) renders "
        "bootstrap confidence intervals unreliable, as the resampling distribution lacks "
        "sufficient support. We conclude that meaningful strategy comparison requires either "
        "substantially more seeds or a larger pool that permits differentiation at maximum budget."
    )

    add_heading_numbered(doc, "7.3", "On-Device Latency", level=2)

    add_body(doc,
        "We profile both model versions on an iPhone 15 Pro Max equipped with the Apple A17 Pro "
        "chip. Table 3 reports latency statistics across three profiling sessions. Session 1 "
        "measures v0 under cold-start conditions, Session 2 measures v0 with the thermal profiling "
        "instrument attached (warm start), and Session 3 measures v1 under standard conditions. "
        "Across all sessions, median latency ranges from 14.0 to 14.5 ms, with 95th-percentile "
        "latency between 22.0 and 26.2 ms. These values represent a 28\u00d7 margin below the "
        "400 ms p50 threshold and a 23\u201327\u00d7 margin below the 600 ms p95 threshold."
    )

    add_table_with_borders(doc,
        headers=["Session", "Model", "Condition", "p50 (ms)", "p95 (ms)", "Mean (ms)", "Threshold"],
        rows=[
            ["1", "v0", "Cold start", "14.0", "26.2", "18.0", "PASS"],
            ["2", "v0", "Warm, TP attached", "14.5", "22.0", "16.8", "PASS"],
            ["3", "v1", "Standard", "14.5", "24.6", "21.3", "PASS"],
        ],
        caption="Physical-device inference latency on iPhone 15 Pro Max (A17 Pro).",
        caption_number=3
    )

    add_body(doc,
        "A noteworthy finding is that physical-device inference is approximately 7\u00d7 faster than "
        "simulator-based measurement, underscoring the importance of on-device profiling. The "
        "transition from v0 to v1\u2014which increases the output head from 24 to 128 classes and "
        "adds approximately 7,000 parameters\u2014increases median latency by only 0.5 ms, "
        "demonstrating that the architectural scaling has negligible impact on inference speed "
        "at this model size."
    )

    add_heading_numbered(doc, "7.4", "Pareto View", level=2)

    add_body(doc,
        "Table 4 presents a Pareto analysis of quality versus efficiency across all model "
        "variants. The tiny_multimodal_v1 model strictly dominates both v0 variants: it achieves "
        "higher test EM (27.5% vs. 10%) with effectively identical latency (14.5 vs. 14.0 ms) "
        "and model size (0.5 MB). The heuristic baseline, while achieving the same test EM as "
        "v0, has no meaningful latency measurement since it performs string lookup rather than "
        "neural inference."
    )

    add_table_with_borders(doc,
        headers=["Model", "Dataset", "Test EM", "Latency p50", "Size"],
        rows=[
            ["question_lookup_v0", "v1", "0.100", "\u2014", "0.1 MB"],
            ["tiny_multimodal_v0", "v1", "0.100", "14.0 ms", "0.5 MB"],
            ["tiny_multimodal_v1", "v2", "0.275", "14.5 ms", "0.5 MB"],
        ],
        caption="Pareto view of quality versus efficiency.",
        caption_number=4
    )

    add_heading_numbered(doc, "7.5", "Effectiveness Scorecard", level=2)

    add_body(doc,
        "Table 5 evaluates both model versions against the six effectiveness thresholds defined "
        "in Section 1. Three of six thresholds are met by both versions: latency p50, latency "
        "p95, and model size. The quality threshold (EM \u2265 70%) remains unmet, with v1 "
        "achieving 27.5%\u2014a substantial improvement over v0's 10% but still 42.5 percentage "
        "points short. Energy and memory metrics are marked as UNAVAILABLE because the current "
        "profiling infrastructure does not yet capture battery drain or peak memory allocation. "
        "Completing these measurements is a priority for future work."
    )

    add_table_with_borders(doc,
        headers=["Metric", "Threshold", "v0", "v1", "Status"],
        rows=[
            ["EM", "\u226570%", "10%", "27.5%", "FAIL"],
            ["Latency p50", "\u2264400 ms", "14.0 ms", "14.5 ms", "PASS"],
            ["Latency p95", "\u2264600 ms", "26.2 ms", "24.6 ms", "PASS"],
            ["Energy", "<5%/hr", "\u2014", "\u2014", "UNAVAILABLE"],
            ["Memory", "<500 MB", "\u2014", "\u2014", "UNAVAILABLE"],
            ["Size", "<100 MB", "96 KB", "0.5 MB", "PASS"],
        ],
        caption="Effectiveness scorecard against six deployment thresholds.",
        caption_number=5
    )

    # ── Section 8: Discussion ────────────────────────────────────────────
    add_heading_numbered(doc, "8", "Discussion")

    add_body(doc,
        "The primary contribution of AXIOM-Mobile at its current stage is infrastructure "
        "validation accompanied by preliminary scaling evidence, rather than a definitive "
        "answer to the research question. We have demonstrated five concrete achievements: "
        "(1) an end-to-end reproducible pipeline spanning dataset creation, model training, "
        "Core ML export, on-device inference, and statistical analysis; (2) an honest status "
        "tracking system that transparently reports which metrics are met, unmet, or unavailable "
        "rather than selectively presenting favorable results; (3) a statistical framework with "
        "bootstrap confidence intervals, paired comparisons, and power-law fitting; (4) physical-"
        "device latency profiling that reveals a 28\u00d7 margin below quality-of-service thresholds; "
        "and (5) preliminary scaling evidence showing that increasing the dataset from 52 to 452 "
        "examples improves test EM from 10% to 27.5%."
    )

    add_body(doc,
        "The key insight emerging from our experiments is that dataset scale\u2014not model "
        "architecture\u2014is the binding constraint for tiny models operating on domain-specific "
        "VQA tasks. The v0-to-v1 transition changes only the training data and output head "
        "dimension while preserving the core architecture, yet test EM improves by 2.75\u00d7. "
        "Conversely, the choice of selection strategy has no detectable effect in the v0 regime, "
        "suggesting that at very small pool sizes the variance introduced by data ordering "
        "dominates any systematic advantage of intelligent selection."
    )

    add_body(doc,
        "An encouraging secondary finding is the negligible latency cost of scaling from 24 "
        "to 128 output classes. The 0.5 ms increase in median latency (14.0 to 14.5 ms) "
        "indicates that the classification head contributes minimally to total inference time, "
        "and that future increases in vocabulary size are unlikely to compromise latency targets."
    )

    add_body(doc,
        "The path forward is clear along several dimensions. First, conducting the full "
        "selection-strategy sweep on dataset v2 will reveal whether intelligent selection "
        "becomes beneficial at larger pool sizes. Second, introducing a stronger model backbone"
        "\u2014such as a LoRA-adapted vision-language model\u2014may unlock substantially higher "
        "accuracy while remaining within mobile resource budgets. Third, instrumenting energy "
        "and memory traces will complete the effectiveness scorecard. Fourth, increasing the "
        "number of seeds from three to at least ten will provide the statistical power needed "
        "for reliable strategy comparisons and scaling-law estimation."
    )

    # ── Section 9: Limitations and Threats to Validity ───────────────────
    add_heading_numbered(doc, "9", "Limitations and Threats to Validity")

    add_body(doc,
        "We identify several significant limitations that constrain the strength of our "
        "conclusions. The most prominent is the quality gap: the best model achieves 27.5% "
        "test EM against a 70% target, leaving substantial room for improvement. No selection-"
        "strategy sweep has been conducted on dataset v2, meaning that our learning curves and "
        "strategy comparisons reflect only the 52-example regime. The model architecture, at "
        "47K parameters with no pretrained backbone, is likely too small to learn robust visual "
        "features for the task\u2014evidenced by the flat learning curves in the v0 regime."
    )

    add_body(doc,
        "Statistical validity is limited by the use of only three seeds per experimental "
        "condition. Bootstrap confidence intervals computed from three observations cannot "
        "capture the true sampling distribution, and our paired comparisons consequently lack "
        "statistical power. A single annotator produced all ground-truth labels, introducing "
        "potential systematic biases that would be detectable only through a dual-annotator "
        "protocol with inter-rater reliability metrics such as Cohen's kappa."
    )

    add_body(doc,
        "Several planned components remain unimplemented: the knowledge-graph-guided selection "
        "strategy, energy and memory profiling, ablation studies to isolate the contribution "
        "of each architectural component, cross-validation for more robust accuracy estimates, "
        "and comparison to production-grade vision-language models that would contextualize "
        "our results against the current state of the art."
    )

    # ── Section 10: Conclusion ───────────────────────────────────────────
    add_heading_numbered(doc, "10", "Conclusion")

    add_body(doc,
        "We have presented AXIOM-Mobile, a fully functional system comprising an iOS testbed "
        "and Python analysis pipeline for measuring minimal training data requirements for "
        "on-device domain reasoning. The system implements three data selection strategies "
        "(random, uncertainty, and diversity), supports reproducible benchmarking through "
        "deterministic input providers and CSV export, and performs honest statistical analysis "
        "that transparently reports both achievements and shortfalls."
    )

    add_body(doc,
        "Our experiments span two regimes: v0, with 52 training examples and 10% test exact "
        "match, and v1, with 452 examples and 27.5% test exact match. The 2.75\u00d7 improvement "
        "confirms dataset scaling as the primary lever for quality improvement in the tiny-model "
        "regime. Physical-device profiling on an iPhone 15 Pro Max demonstrates 14.0\u201314.5 ms "
        "median inference latency, providing a 28\u00d7 margin below the 400 ms quality-of-service "
        "threshold. Three of six effectiveness thresholds are met, with the quality target "
        "(EM \u2265 70%) remaining the principal outstanding challenge."
    )

    add_body(doc,
        "The research question\u2014determining k*\u2014remains open. Our results suggest that k* "
        "lies substantially above 452 examples for the current model architecture, and that "
        "either a larger dataset, a stronger model backbone, or both will be required to reach "
        "the 70% EM threshold. The infrastructure we have built is complete and ready to support "
        "these investigations. Future work will focus on conducting v2 sweeps, integrating a "
        "LoRA-adapted vision-language model, completing energy and memory profiling, and "
        "increasing experimental seeds to enable reliable statistical inference."
    )

    # ── Acknowledgments ──────────────────────────────────────────────────
    add_heading_numbered(doc, "", "Acknowledgments")
    # Override to remove number
    doc.paragraphs[-1].runs[0].text = "Acknowledgments"

    add_body(doc,
        "This work was conducted as part of CMPT 416 at Simon Fraser University. We thank "
        "the course instructor for guidance on experimental methodology, honest reporting "
        "practices, and the importance of infrastructure validation in empirical machine "
        "learning research."
    )

    # ── References ───────────────────────────────────────────────────────
    add_heading_numbered(doc, "", "References")
    doc.paragraphs[-1].runs[0].text = "References"

    references = [
        '[1] Sener, O. and Savarese, S. "Active Learning for Convolutional Neural Networks: A Core-Set Approach." In Proceedings of the International Conference on Learning Representations (ICLR), 2018.',
        '[2] Ash, J.T., Zhang, C., Krishnamurthy, A., Langford, J., and Agarwal, A. "Deep Batch Active Learning by Diverse, Uncertain Gradient Lower Bounds." In Proceedings of the International Conference on Learning Representations (ICLR), 2020.',
        '[3] Antol, S., Agrawal, A., Lu, J., Mitchell, M., Batra, D., Zitnick, C.L., and Parikh, D. "VQA: Visual Question Answering." In Proceedings of the IEEE International Conference on Computer Vision (ICCV), 2015.',
        '[4] Kaplan, J., McCandlish, S., Henighan, T., Brown, T.B., Chess, B., Child, R., Gray, S., Radford, A., Wu, J., and Amodei, D. "Scaling Laws for Neural Language Models." arXiv preprint arXiv:2001.08361, 2020.',
        '[5] Apple. "Core ML Documentation." developer.apple.com/documentation/coreml, 2023.',
        '[6] Ignatov, A., Timofte, R., Chou, W., Wang, K., Wu, M., Hartley, T., and Van Gool, L. "AI Benchmark: All About Deep Learning on Smartphones." In ICCV Workshop on Advances in Mobile Application Development, 2019.',
        '[7] Hestness, J., Narang, S., Ardalani, N., Diamos, G., Jun, H., Kianinejad, H., Patwary, M., Yang, Y., and Zhou, Y. "Deep Learning Scaling is Predictable, Empirically." arXiv preprint arXiv:1712.00409, 2017.',
        '[8] Settles, B. "Active Learning Literature Survey." Computer Sciences Technical Report 1648, University of Wisconsin\u2013Madison, 2009.',
        '[9] Howard, A., Sandler, M., Chu, G., Chen, L.C., Chen, B., Tan, M., Wang, W., Zhu, Y., Pang, R., Vasudevan, V., Le, Q.V., and Adam, H. "Searching for MobileNetV3." In Proceedings of the IEEE International Conference on Computer Vision (ICCV), 2019.',
        '[10] Liu, H., Li, C., Wu, Q., and Lee, Y.J. "Visual Instruction Tuning." In Advances in Neural Information Processing Systems (NeurIPS), 2023.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run = p.add_run(ref)
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    # ── Page numbers ─────────────────────────────────────────────────────
    add_page_numbers(doc)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"Paper saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate()
